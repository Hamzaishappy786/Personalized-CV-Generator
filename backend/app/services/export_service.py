from pathlib import Path
from typing import Literal

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

from app.schemas.cv_schema import CVData
from app.services.model_service import infer_cv_data
from app.services.template_service import get_template


EXPORT_DIR = Path("storage/exports")


def _ensure_export_dir() -> Path:
    EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    return EXPORT_DIR


def _build_lines(cv_data: CVData, template_id: str) -> list[str]:
    template = get_template(template_id)
    lines = [
        cv_data.name or "Unnamed Candidate",
        template["name"],
        "",
        f"Summary: {cv_data.summary or ''}",
    ]
    if cv_data.skills:
        lines += ["", "Skills: " + ", ".join(cv_data.skills)]
    return lines


def _export_pdf(path: Path, cv_data: CVData, template_id: str) -> None:
    c = canvas.Canvas(str(path), pagesize=A4)
    width, height = A4
    y = height - 72
    for line in _build_lines(cv_data, template_id):
        c.drawString(72, y, line)
        y -= 16
        if y < 72:
            c.showPage()
            y = height - 72
    c.save()


def _export_docx(path: Path, cv_data: CVData, template_id: str) -> None:
    document = Document()
    template = get_template(template_id)
    document.add_heading(cv_data.name or "Unnamed Candidate", level=1)
    document.add_paragraph(template["name"])
    document.add_paragraph(cv_data.summary or "")
    if cv_data.skills:
        document.add_heading("Skills", level=2)
        document.add_paragraph(", ".join(cv_data.skills))
    document.save(path)


def export_document(text: str, template_id: str, format_name: Literal["pdf", "docx"]) -> dict:
    cv_data = infer_cv_data(text)
    export_dir = _ensure_export_dir()
    filename = f"cv-export-{template_id}.{format_name}"
    output_path = export_dir / filename

    if format_name == "docx":
        _export_docx(output_path, cv_data, template_id)
    else:
        _export_pdf(output_path, cv_data, template_id)

    return {"format": format_name, "template_id": template_id, "path": str(output_path)}
